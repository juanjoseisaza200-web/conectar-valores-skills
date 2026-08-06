"""
vf3_styles.py — Reproduce el ADN visual EXACTO del Informe Peers Aseo VF3.

Validado contra `assets/vf3_dna/*.json` (extraído por inspección XML directa).
NO modificar constantes ni props XML sin re-extraer del VF3 original.

USO: from peers-format-applier.templates.vf3_styles import *
     doc = open_blank_template()
     add_h1(doc, "Sección X")
     add_paragraph(doc, "...")
     add_table_vf3(doc, headers=[...], rows=[...], highlight_row=2)
     doc.save("output.docx")
"""

import os, shutil, copy
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =====================================================================
# PALETA HEX VF3 — única autorizada (verificada por inventario XML)
# =====================================================================
NAVY     = "17375E"   # 399 runs en VF3 + 25 bordes
DORADO   = "C9A449"   # 154 bordes (header tabla bottom)
HAIRLINE = "BFBFBF"   # 1,055 bordes (filas cuerpo bottom)
MINT     = "E7F0EF"   # 76 fills (fila destacada — único fill)
TEAL     = "5D9E9D"   # 1 run (decorativo aislado)
BLACK    = "000000"   # Cuerpo (heredado, no se setea)
WHITE    = "FFFFFF"

# Colores PROHIBIDOS (no aparecen en VF3 — generador NO debe usarlos)
FORBIDDEN_COLORS = {"EDEDED", "4472C4", "70AD47", "FF0000", "E67E22", "27AE60"}

# =====================================================================
# TAMAÑOS REALES (verificados en font_inventory.json)
# =====================================================================
# En XML: <w:sz val="N"/> donde N = puntos × 2 (half-points)
SZ_BODY        = "18"  # 9pt   — dominante: 1,252 runs
SZ_TABLE_HEAD  = "19"  # 9.5pt — 156 runs (header tablas)
SZ_FOOTNOTE    = "17"  # 8.5pt — 55 runs (captions, notas)
SZ_H1          = "24"  # 12pt  — 30 runs (sección)
SZ_H2          = "20"  # 10pt
SZ_H3          = "19"  # 9.5pt

# Sizes de bordes (en eighths of a point)
BORDER_DORADO_SZ   = "12"  # 1.5pt — header bottom
BORDER_HAIRLINE_SZ = "2"   # 0.25pt — fila cuerpo bottom

# =====================================================================
# RUTAS
# =====================================================================
HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.normpath(os.path.join(HERE, '..', 'assets'))
BLANK_TEMPLATE = os.path.join(ASSETS, 'vf3_blank_template.docx')


# =====================================================================
# CREAR DOC DESDE TEMPLATE CLONADO
# =====================================================================
def open_blank_template():
    """Abre vf3_blank_template.docx (clon del VF3 con cuerpo vacío).

    Hereda automáticamente:
    - 168 estilos del VF3 original
    - Margenes exactos (1.20" / 0.79" / 0.79" / 0.79")
    - Header / Footer
    - Theme con paleta CV
    - settings.xml + fontTable.xml

    Retorna python-docx Document listo para agregar contenido.
    """
    if not os.path.exists(BLANK_TEMPLATE):
        raise FileNotFoundError(f"Template no encontrado: {BLANK_TEMPLATE}. Correr build_blank_template.py")
    return Document(BLANK_TEMPLATE)


# =====================================================================
# HELPERS XML BAJO NIVEL
# =====================================================================

def _set_run_props(run, *, size=None, bold=False, italic=False, color_hex=None):
    """Aplica props a un run con XML directo (replicando exactamente VF3)."""
    rPr = run._r.get_or_add_rPr()
    # Forzar Arial en TODOS los slots
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for slot in ('ascii', 'hAnsi', 'cs', 'eastAsia'):
        rFonts.set(qn(f'w:{slot}'), 'Arial')
    if bold:
        b = rPr.find(qn('w:b'))
        if b is None:
            b = OxmlElement('w:b'); rPr.append(b)
    if italic:
        i = rPr.find(qn('w:i'))
        if i is None:
            i = OxmlElement('w:i'); rPr.append(i)
    if size:
        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = OxmlElement('w:sz'); rPr.append(sz)
        sz.set(qn('w:val'), str(size))
    if color_hex:
        color = rPr.find(qn('w:color'))
        if color is None:
            color = OxmlElement('w:color'); rPr.append(color)
        color.set(qn('w:val'), color_hex)


def _set_para_spacing(p, *, beforeLines=None, before=None, afterLines=None, after=None):
    """Pone props de spacing en un párrafo (XML directo)."""
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing'); pPr.append(spacing)
    if beforeLines is not None: spacing.set(qn('w:beforeLines'), str(beforeLines))
    if before is not None:      spacing.set(qn('w:before'),      str(before))
    if afterLines is not None:  spacing.set(qn('w:afterLines'),  str(afterLines))
    if after is not None:       spacing.set(qn('w:after'),       str(after))


def _set_cell_shading(cell, hex_color):
    """Pinta el fill de una celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    # Quitar shd previo si existe
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_borders_vf3(cell, *, top='nil', left='nil', right='nil',
                          bottom_val='nil', bottom_sz=None, bottom_color=None):
    """Pone los bordes exactos estilo VF3 (laterales nil, solo bottom configurable)."""
    tcPr = cell._tc.get_or_add_tcPr()
    # Eliminar tcBorders previo
    for b in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(b)
    tcBorders = OxmlElement('w:tcBorders')

    # top, left, right = nil siempre
    for side in ('top', 'left', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        tcBorders.append(el)

    # bottom: configurable
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), bottom_val)
    if bottom_val != 'nil':
        bot.set(qn('w:sz'), str(bottom_sz))
        bot.set(qn('w:space'), '0')
        bot.set(qn('w:color'), bottom_color)
    tcBorders.append(bot)

    tcPr.append(tcBorders)


def _set_cell_valign(cell, val='center'):
    tcPr = cell._tc.get_or_add_tcPr()
    for v in tcPr.findall(qn('w:vAlign')):
        tcPr.remove(v)
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), val)
    tcPr.append(vAlign)


# =====================================================================
# HEADINGS y PÁRRAFOS
# =====================================================================

def add_h1(doc, text):
    """Sección — Heading 1 (12pt bold Navy centrado)."""
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    _set_run_props(r, size=SZ_H1, bold=True, color_hex=NAVY)
    return p


def add_h2(doc, text):
    """Sub-sección — Heading 2 (10pt bold Navy izquierda)."""
    p = doc.add_paragraph(style='Heading 2')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _set_run_props(r, size=SZ_H2, bold=True, color_hex=NAVY)
    return p


def add_h3(doc, text):
    """Sub-sub — Heading 3 (9.5pt bold Navy izquierda)."""
    p = doc.add_paragraph(style='Heading 3')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _set_run_props(r, size=SZ_H3, bold=True, color_hex=NAVY)
    return p


def add_paragraph(doc, text, justify=True):
    """Párrafo cuerpo (9pt Arial, justificado por defecto)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _set_run_props(r, size=SZ_BODY)
    return p


def add_bullet(doc, text):
    """Bullet justificado (9pt Arial)."""
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    if p.runs:
        for r in list(p.runs): r.text = ""
    r = p.add_run(text)
    _set_run_props(r, size=SZ_BODY)
    return p


def add_caption(doc, text):
    """Caption sobre tabla/figura (Navy bold 8.5pt centrado)."""
    p = doc.add_paragraph(style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Quitar prefijo "Tabla 1: " automático del estilo Caption si python-docx lo agrega
    for r in list(p.runs):
        r.text = ""
    r = p.add_run(text)
    _set_run_props(r, size=SZ_FOOTNOTE, bold=True, color_hex=NAVY)
    return p


def add_footnote(doc, text):
    """Nota al pie de tabla (Slate italic 8.5pt izquierda)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _set_run_props(r, size=SZ_FOOTNOTE, italic=True, color_hex="4A5568")
    return p


# =====================================================================
# FORMULA BLOCK — bloque de fórmula/cálculo SIN caja (párrafos)
# =====================================================================

def add_formula_block(doc, text, *, title=None):
    """Bloque de fórmula/cálculo renderizado como párrafos limpios — SIN caja.

    Uso correcto para fórmulas matemáticas, cálculos intermedios o resúmenes
    estructurados que aparecen como "tablas 1x1" en el doc original. NO usa
    fill ni borde — son párrafos con título bold Navy + cuerpo en sangría leve.

    Args:
        doc: python-docx Document
        text: contenido (puede tener saltos de línea)
        title: encabezado opcional (bold Navy, sin sangría)
    """
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_para_spacing(p, before=160, after=40)
        r = p.add_run(title)
        _set_run_props(r, size=SZ_BODY, bold=True, color_hex=NAVY)

    lines = [ln.strip() for ln in str(text).split('\n') if ln.strip()]
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_para_spacing(p, before=20, after=20)
        p.paragraph_format.left_indent = Inches(0.20)  # sangría leve
        r = p.add_run(line)
        _set_run_props(r, size=SZ_BODY)

    # Pequeño cierre vertical
    p = doc.add_paragraph()
    _set_para_spacing(p, before=20, after=80)


# =====================================================================
# CALLOUT BOX — caja destacada Mint (USO RESERVADO — no para fórmulas)
# =====================================================================
# Usar SOLO cuando el usuario explícitamente quiera una caja destacada.
# Para fórmulas y cálculos del doc original, usar add_formula_block.

def add_callout(doc, text, *, title=None):
    """Caja destacada con fill Mint, borde inferior dorado, texto Navy.

    Para reemplazar tablas 1x1 originales que son cajas de fórmula/cálculo.
    Cada salto de línea del texto se renderiza como un sub-párrafo.
    """
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)

    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); table._tbl.insert(0, tblPr)
    jc = tblPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc'); tblPr.append(jc)
    jc.set(qn('w:val'), 'center')

    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    cell.text = ""
    _set_cell_shading(cell, MINT)
    _set_cell_borders_vf3(cell,
        bottom_val='single', bottom_sz=BORDER_DORADO_SZ, bottom_color=DORADO)
    _set_cell_valign(cell, 'center')

    # Limpiar párrafo inicial vacío
    cell._tc.remove(cell.paragraphs[0]._p)

    if title:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, before=80, after=40)
        r = p.add_run(title)
        _set_run_props(r, size=SZ_TABLE_HEAD, bold=True, color_hex=NAVY)

    lines = [ln.strip() for ln in str(text).split('\n') if ln.strip()]
    for i, line in enumerate(lines):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if i == 0 and not title:
            _set_para_spacing(p, before=100, after=40)
        elif i == len(lines) - 1:
            _set_para_spacing(p, before=20, after=100)
        else:
            _set_para_spacing(p, before=20, after=20)
        r = p.add_run(line)
        _set_run_props(r, size=SZ_BODY, color_hex=NAVY)

    return table


# =====================================================================
# IMAGEN / FIGURA — inserta imagen con estilo CV
# =====================================================================

def add_image(doc, image_path, *, width_in=6.5):
    """Inserta imagen centrada escalada a width_in pulgadas."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(p, before=120, after=60)
    run = p.add_run()
    shape = run.add_picture(image_path, width=Inches(width_in))
    return shape


def add_figure_with_caption(doc, image_path, caption_text, *, width_in=6.5, source=None):
    """Inserta figura + caption Navy bold debajo + opcional fuente italic Slate."""
    add_image(doc, image_path, width_in=width_in)
    add_caption(doc, caption_text)
    if source:
        add_footnote(doc, source)


# =====================================================================
# TABLA VF3 — replica EXACTA del ADN extraído
# =====================================================================

def add_table_vf3(doc, headers, rows, *, highlight_row=None, col_widths_in=None):
    """Tabla con estilo VF3 EXACTO:
       - Tabla centrada en página (jc=center), tblLayout fixed
       - Sin bordes externos
       - Header: blanco, texto Navy bold 9.5pt centrado, bottom dorado 1.5pt
       - Cuerpo: blanco, texto negro 9pt regular centrado, bottom hairline 0.25pt
       - Fila destacada: fill Mint, texto Navy bold 9pt centrado
       - Sin bandas alternadas
       - Sin filas Total con color especial (si se quiere total, simplemente bold)

    Args:
        doc: python-docx Document
        headers: lista de strings (textos del header)
        rows: lista de listas (cada sublista = una fila)
        highlight_row: índice 0-based de la fila a destacar en Mint (típico: la del cliente)
        col_widths_in: lista de anchos en pulgadas; si None, distribuye uniforme sobre 6.92"
    """
    n_cols = len(headers)
    n_rows = len(rows) + 1  # +1 por header

    # Crear tabla en estilo "Table Grid" base (luego sobreescribimos bordes)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.autofit = False

    # tblPr: centrar tabla en página, layout fixed
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); table._tbl.insert(0, tblPr)
    # tblW auto
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    # jc center
    jc = tblPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc'); tblPr.append(jc)
    jc.set(qn('w:val'), 'center')
    # tblLayout fixed
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout'); tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')
    # tblLook
    tblLook = tblPr.find(qn('w:tblLook'))
    if tblLook is None:
        tblLook = OxmlElement('w:tblLook'); tblPr.append(tblLook)
    tblLook.set(qn('w:val'), '04A0')
    tblLook.set(qn('w:firstRow'), '1')
    tblLook.set(qn('w:lastRow'), '0')
    tblLook.set(qn('w:firstColumn'), '1')
    tblLook.set(qn('w:lastColumn'), '0')
    tblLook.set(qn('w:noHBand'), '0')
    tblLook.set(qn('w:noVBand'), '1')

    # Eliminar tblBorders si vino del estilo base
    for tb in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(tb)

    # Anchos de columna
    if col_widths_in is None:
        # Distribuir uniforme sobre 6.92"
        col_widths_in = [6.92 / n_cols] * n_cols
    for i, w in enumerate(col_widths_in):
        for row in table.rows:
            row.cells[i].width = Inches(w)

    # ===== HEADER (fila 0) =====
    for i, txt in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        _set_cell_borders_vf3(cell,
            bottom_val='single', bottom_sz=BORDER_DORADO_SZ, bottom_color=DORADO)
        _set_cell_valign(cell, 'center')
        # NO shd (fondo blanco)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, beforeLines=20, before=48, afterLines=20, after=48)
        r = p.add_run(str(txt))
        _set_run_props(r, size=SZ_TABLE_HEAD, bold=True, color_hex=NAVY)

    # Marcar primera fila como tblHeader (se repite si se parte)
    tr0 = table.rows[0]._tr
    trPr = tr0.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr'); tr0.insert(0, trPr)
    if trPr.find(qn('w:tblHeader')) is None:
        trPr.append(OxmlElement('w:tblHeader'))
    jc_tr = trPr.find(qn('w:jc'))
    if jc_tr is None:
        jc_tr = OxmlElement('w:jc'); trPr.append(jc_tr)
    jc_tr.set(qn('w:val'), 'center')

    # ===== CUERPO =====
    for r_idx, row_data in enumerate(rows):
        is_highlight = (highlight_row is not None and r_idx == highlight_row)
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            _set_cell_borders_vf3(cell,
                bottom_val='single', bottom_sz=BORDER_HAIRLINE_SZ, bottom_color=HAIRLINE)
            _set_cell_valign(cell, 'center')

            if is_highlight:
                _set_cell_shading(cell, MINT)

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if is_highlight:
                # Spacing simplificado en fila destacada (igual que VF3)
                _set_para_spacing(p, before=20, after=20)
            else:
                _set_para_spacing(p, beforeLines=20, before=48, afterLines=20, after=48)

            r = p.add_run(str(val))
            if is_highlight:
                _set_run_props(r, size=SZ_BODY, bold=True, color_hex=NAVY)
            else:
                _set_run_props(r, size=SZ_BODY)

        # trPr center
        tr_now = table.rows[r_idx + 1]._tr
        trPr = tr_now.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr'); tr_now.insert(0, trPr)
        jc_tr = trPr.find(qn('w:jc'))
        if jc_tr is None:
            jc_tr = OxmlElement('w:jc'); trPr.append(jc_tr)
        jc_tr.set(qn('w:val'), 'center')

    return table


# =====================================================================
# SMOKE TEST
# =====================================================================
if __name__ == "__main__":
    doc = open_blank_template()
    add_h1(doc, "Smoke Test — VF3 styles aplicados desde template")
    add_paragraph(doc,
        "Este documento se construyó clonando vf3_blank_template.docx y agregando contenido "
        "con las funciones de vf3_styles.py. Heredó automáticamente los 168 estilos del VF3 "
        "original, los márgenes Letter 1.20\"/0.79\", la fuente Arial, la paleta hex, "
        "el footer y el theme.")

    add_h2(doc, "Tabla de prueba — peers ficticios")
    add_caption(doc, "Tabla 1 — Universo de prueba")
    add_table_vf3(doc,
        headers=["#", "Operador", "MW", "Ingresos COP MM", "ROE"],
        rows=[
            ["1", "Empresa A",  "1,200", "850,000", "14.9%"],
            ["2", "Empresa B",  "980",   "720,500", "12.6%"],
            ["3", "Cliente",    "1,450", "920,300", "16.8%"],
            ["4", "Empresa D",  "560",   "410,200", "13.7%"],
            ["5", "Empresa E",  "720",   "540,100", "11.4%"],
        ],
        highlight_row=2,
        col_widths_in=[0.5, 1.7, 1.0, 2.0, 1.0],
    )
    add_footnote(doc, "Fuente: smoke test — peers ficticios. Cifras en COP MM nominales.")

    out = os.path.join(ASSETS, "smoke_test_vf3.docx")
    doc.save(out)
    print(f"[OK] Smoke test guardado: {out}")
