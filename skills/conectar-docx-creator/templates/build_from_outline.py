"""
build_from_outline.py — Construye un .docx desde un dict de especificación.

USO:
    from build_from_outline import build_document
    build_document(doc_spec)  # dict con keys: output_path, portada, secciones

doc_spec schema:
    {
        "output_path": "C:\\...\\nombre.docx",
        "portada": {                                 # opcional
            "empresa":       str,
            "titulo":        str,
            "subtitulo":     str,                    # opcional
            "descripcion":   str,                    # opcional, italic
            "fecha":         str,
            "preparador":    str (default "CONECTAR VALORES S.A.S."),
            "confidencial":  bool (default True),
        },
        "secciones": [
            {"type": "h1", "text": "..."},
            {"type": "h2", "text": "..."},
            {"type": "h3", "text": "..."},
            {"type": "paragraph", "text": "..."},
            {"type": "bullets", "items": ["...", ...]},
            {"type": "caption", "text": "Tabla N — ..."},
            {"type": "table",
                "headers": [...], "rows": [[...], ...],
                "highlight_row": int|None,
                "col_widths_in": [float, ...] | None,
                "footnote": str | None
            },
            {"type": "footnote", "text": "..."},
            {"type": "page_break"},
        ]
    }
"""
import os, sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
from vf3_styles import (
    open_blank_template, add_h1, add_h2, add_h3, add_paragraph, add_bullet,
    add_caption, add_footnote, add_table_vf3, add_callout, add_image, add_figure_with_caption,
    _set_run_props, _set_para_spacing,
    SZ_BODY, SZ_FOOTNOTE, NAVY, Inches,
)
from vf3_charts import (
    chart_single_bar, chart_double_bar, chart_line, chart_pie,
)
import tempfile
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_portada(doc, p):
    """Construye portada estándar VF3."""
    # Empresa (16pt bold Navy centrado)
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(pp, before=2400)
    r = pp.add_run(p.get("empresa", ""))
    _set_run_props(r, size="32", bold=True, color_hex=NAVY)

    # Título (13pt bold Navy)
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(pp, before=240)
    r = pp.add_run(p.get("titulo", ""))
    _set_run_props(r, size="26", bold=True, color_hex=NAVY)

    # Subtítulo (11pt bold Navy)
    if p.get("subtitulo"):
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(pp, before=120)
        r = pp.add_run(p["subtitulo"])
        _set_run_props(r, size="22", bold=True, color_hex=NAVY)

    # Descripción (italic 9pt centrado)
    if p.get("descripcion"):
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(pp, before=480)
        r = pp.add_run(p["descripcion"])
        _set_run_props(r, size=SZ_BODY, italic=True)

    # Fecha (11pt bold Navy)
    if p.get("fecha"):
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(pp, before=1200)
        r = pp.add_run(p["fecha"])
        _set_run_props(r, size="22", bold=True, color_hex=NAVY)

    # Preparador (11pt bold Navy)
    preparador = p.get("preparador", "CONECTAR VALORES S.A.S.")
    pp = doc.add_paragraph()
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(pp, before=2400)
    r = pp.add_run(preparador)
    _set_run_props(r, size="22", bold=True, color_hex=NAVY)

    # Confidencial
    if p.get("confidencial", True):
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(pp, before=120)
        r = pp.add_run("Documento Confidencial · Para circulación interna")
        _set_run_props(r, size=SZ_FOOTNOTE, italic=True, color_hex="4A5568")

    doc.add_page_break()


def build_document(spec):
    """Construye .docx desde dict de spec.

    Args:
        spec: dict con keys output_path, portada (opcional), secciones (lista)
    """
    output_path = spec["output_path"]
    folder = os.path.dirname(output_path)
    if folder and not os.path.exists(folder):
        os.makedirs(folder, exist_ok=True)

    doc = open_blank_template()

    # Portada
    if "portada" in spec:
        print("[INFO] Construyendo portada...")
        _add_portada(doc, spec["portada"])

    # Secciones
    n_tables = 0
    for i, item in enumerate(spec.get("secciones", [])):
        t = item.get("type")
        if t == "h1":
            add_h1(doc, item["text"])
        elif t == "h2":
            add_h2(doc, item["text"])
        elif t == "h3":
            add_h3(doc, item["text"])
        elif t == "paragraph":
            add_paragraph(doc, item["text"])
        elif t == "bullets":
            for b in item.get("items", []):
                add_bullet(doc, b)
        elif t == "caption":
            add_caption(doc, item["text"])
        elif t == "table":
            n_tables += 1
            headers = item["headers"]
            rows = item["rows"]
            highlight = item.get("highlight_row")
            widths = item.get("col_widths_in")
            add_table_vf3(doc,
                headers=headers, rows=rows,
                highlight_row=highlight,
                col_widths_in=widths)
            print(f"  Tabla {n_tables}: {len(headers)} cols x {len(rows)} filas")
            if item.get("footnote"):
                add_footnote(doc, item["footnote"])
        elif t == "footnote":
            add_footnote(doc, item["text"])
        elif t == "page_break":
            doc.add_page_break()
        elif t == "callout":
            add_callout(doc, item["text"], title=item.get("title"))
        elif t == "image":
            add_image(doc, item["path"], width_in=item.get("width_in", 6.5))
            if item.get("caption"):
                add_caption(doc, item["caption"])
            if item.get("source"):
                add_footnote(doc, item["source"])
        elif t == "chart":
            # Generar chart PNG en tempdir + insertarlo
            kind = item.get("kind", "single_bar")
            tmp_png = tempfile.NamedTemporaryFile(delete=False, suffix='.png').name
            if kind == "single_bar":
                chart_single_bar(
                    title=item["title"], categories=item["categories"],
                    values=item["values"], ylabel=item.get("ylabel"),
                    source=item.get("source"), output_path=tmp_png,
                    fmt_pct=item.get("fmt_pct", True),
                )
            elif kind == "double_bar":
                chart_double_bar(
                    title=item["title"], left=tuple(item["left"]), right=tuple(item["right"]),
                    left_ylabel=item.get("left_ylabel"),
                    right_ylabel=item.get("right_ylabel"),
                    source=item.get("source"), output_path=tmp_png,
                )
            elif kind == "line":
                chart_line(
                    title=item["title"], x_categories=item["x_categories"],
                    series=item["series"], ylabel=item.get("ylabel"),
                    source=item.get("source"), output_path=tmp_png,
                )
            elif kind == "pie":
                chart_pie(
                    title=item["title"], categories=item["categories"],
                    values=item["values"], source=item.get("source"),
                    output_path=tmp_png,
                )
            else:
                print(f"[WARN] kind chart desconocido: {kind}")
                continue
            add_image(doc, tmp_png, width_in=item.get("width_in", 6.5))
            if item.get("caption"):
                add_caption(doc, item["caption"])
            print(f"  Chart ({kind}): {item.get('title','')[:60]}")
        else:
            print(f"[WARN] Tipo desconocido en item {i}: {t}")

    doc.save(output_path)
    print(f"\n[OK] Guardado: {output_path}")
    print(f"     Tamaño: {os.path.getsize(output_path)/1024:.1f} KB")
    return output_path


# Smoke test
if __name__ == "__main__":
    spec_demo = {
        "output_path": os.path.join(
            os.path.dirname(os.path.abspath(__file__)),
            "..", "examples", "smoke_creator.docx"
        ),
        "portada": {
            "empresa": "EJEMPLO S.A.S",
            "titulo": "MEMO DE PRUEBA",
            "subtitulo": "SMOKE TEST DEL CREATOR",
            "descripcion": "Documento de prueba para validar el constructor desde outline.",
            "fecha": "Mayo 2026",
            "confidencial": True,
        },
        "secciones": [
            {"type": "h1", "text": "1. Contexto"},
            {"type": "paragraph",
             "text": "Este es un memo de prueba para verificar que el constructor "
                     "produce documentos conformes al ADN VF3 de Conectar Valores."},
            {"type": "h2", "text": "1.1 Antecedentes"},
            {"type": "bullets", "items": [
                "Bullet uno con justificación.",
                "Bullet dos para confirmar sangría y centrado.",
                "Bullet tres con contenido más extenso para validar el wrap de líneas en un párrafo justificado típico del cuerpo del documento.",
            ]},
            {"type": "h1", "text": "2. Cifras de ejemplo"},
            {"type": "caption", "text": "Tabla 1 — Universo de prueba"},
            {"type": "table",
             "headers": ["Peer", "Ingresos COP MM", "EBITDA %", "ROE"],
             "rows": [
                ["Empresa A", "850,000", "32.4%", "14.9%"],
                ["Empresa B", "720,500", "28.7%", "12.6%"],
                ["Cliente",   "920,300", "34.1%", "16.8%"],
                ["Empresa D", "410,200", "25.9%", "13.7%"],
             ],
             "highlight_row": 2,
             "col_widths_in": [1.5, 1.8, 1.3, 1.3],
             "footnote": "Fuente: smoke test, peers ficticios."},
            {"type": "h2", "text": "2.1 Lectura"},
            {"type": "paragraph",
             "text": "El cliente exhibe el mayor ROE de la muestra con margen EBITDA "
                     "consistente en el tercio superior del sector."},
        ],
    }
    out = build_document(spec_demo)
    print(f"\nSmoke test creator listo: {out}")
