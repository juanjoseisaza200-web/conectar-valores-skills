"""
reformat_template.py — Plantilla del re-formateador.

Uso:
    python reformat_template.py "<input.docx>" ["<output.docx>"]

Si no se pasa output, se infiere como <input>_CV.docx en la misma carpeta del input.

FIXES iter 2:
- Tablas 1x1 multilínea (cajas fórmula/cálculo) → add_callout (no add_table_vf3)
- Tablas con n_cols == 0 o n_rows == 0 → omitidas (no crear tabla vacía)
- Imágenes embedidas en input → extraídas a tempdir y re-insertadas con add_image
- Tablas con celdas anidadas → degradadas gracefully (toman primera tabla anidada)
"""
import os, sys, io, tempfile, zipfile, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
from vf3_styles import (
    open_blank_template, add_h1, add_h2, add_h3, add_paragraph, add_bullet,
    add_caption, add_footnote, add_table_vf3, add_formula_block, add_image,
    _set_run_props, _set_para_spacing,
    SZ_BODY, SZ_FOOTNOTE, NAVY, Inches,
)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def extract_table(tbl):
    headers = [c.text.strip() for c in tbl.rows[0].cells]
    rows = [[c.text.strip() for c in row.cells] for row in tbl.rows[1:]]
    return headers, rows


def is_formula_block_table(tbl):
    """Detecta si una tabla 1x1 contiene una fórmula/cálculo que debe
    renderizarse como bloque de párrafos (sin caja), NO como tabla.
    """
    if not tbl.rows:
        return False
    n_cols = len(tbl.columns)
    n_rows = len(tbl.rows)
    if n_cols == 1 and n_rows == 1:
        text = tbl.rows[0].cells[0].text
        if '\n' in text or len(text) > 120:
            return True
    return False


def is_empty_table(tbl):
    """Tabla totalmente vacía o sin contenido textual real."""
    if not tbl.rows:
        return True
    if len(tbl.columns) == 0:
        return True
    has_text = False
    for row in tbl.rows:
        for c in row.cells:
            if c.text.strip():
                has_text = True
                break
        if has_text:
            break
    return not has_text


def detect_highlight(headers, rows):
    if 'Severidad' in headers:
        for i, row in enumerate(rows):
            last = (row[-1] or '').upper()
            if 'CRÍTICA' in last or 'CRITICA' in last:
                return i
    for i, row in enumerate(rows):
        first = (row[0] or '').lower()
        if any(kw in first for kw in ['cliente', 'mediana']):
            return i
    return None


def col_widths_for(headers, n_rows):
    n = len(headers)
    if n == 0: return None
    if n == 2: return [1.2, 5.7]
    if n == 3:
        # Si la 1ª columna es texto largo descriptivo
        return [2.0, 2.46, 2.46]
    if n == 4: return [1.5, 1.85, 1.85, 1.72]
    if n == 5 and 'Departamento' in headers: return [1.4, 0.75, 1.25, 1.25, 1.4]
    if n == 5: return [1.85, 1.27, 1.27, 1.27, 1.27]
    if n == 6: return [1.7, 1.0, 1.0, 1.05, 1.05, 1.12]
    if n == 7: return [1.7, 0.87, 0.87, 0.87, 0.87, 0.87, 0.87]
    if n == 8: return [0.6, 1.05, 1.0, 0.7, 0.75, 0.7, 0.87, 1.25]
    return [6.92/n] * n


def formula_title_from_text(text):
    """Si el callout texto es 'TÍTULO / contenido', extrae título y resto."""
    lines = [ln.strip() for ln in text.split('\n') if ln.strip()]
    if not lines:
        return None, text
    # Si la 1ª línea está toda en mayúsculas y es <80 chars, es título
    first = lines[0]
    if first.isupper() and len(first) < 80:
        return first, '\n'.join(lines[1:])
    # Si hay un " / " separador en la 1ª línea
    if ' / ' in first:
        parts = first.split(' / ', 1)
        if parts[0].isupper() or len(parts[0]) < 60:
            return parts[0], parts[1] + '\n' + '\n'.join(lines[1:]) if len(lines) > 1 else parts[1]
    return None, text


def extract_images_from_docx(input_path):
    """Extrae todas las imágenes embedded del .docx a un tempdir.

    Returns: dict {filename: ruta_temp_extraida}
    """
    tmpdir = tempfile.mkdtemp(prefix="vf3_imgs_")
    images = {}
    with zipfile.ZipFile(input_path) as z:
        for name in z.namelist():
            if name.startswith('word/media/'):
                basename = os.path.basename(name)
                out_path = os.path.join(tmpdir, basename)
                with open(out_path, 'wb') as f:
                    f.write(z.read(name))
                images[basename] = out_path
    return images, tmpdir


def get_image_rels(src):
    """Retorna dict {rId: filename} mapeando relationships a archivos media."""
    rels = {}
    for r_id, rel in src.part.rels.items():
        if 'image' in rel.reltype:
            rels[r_id] = os.path.basename(rel.target_ref)
    return rels


def reformat(input_path, output_path):
    print(f"[INFO] Input:  {input_path}")
    print(f"[INFO] Output: {output_path}")

    src = Document(input_path)
    dst = open_blank_template()

    # Extraer imágenes embebidas a tempdir
    image_files, tmpdir = extract_images_from_docx(input_path)
    rels_map = get_image_rels(src)
    print(f"[INFO] Imágenes detectadas en input: {len(image_files)}")

    table_iter = iter(src.tables)
    n_paragraphs_consumed = 0
    n_tables_rendered = 0
    n_callouts = 0
    n_images_rendered = 0
    n_tables_skipped = 0

    for elem in src.element.body.iterchildren():
        tag = elem.tag.split('}')[1]

        if tag == 'p':
            try:
                p = src.paragraphs[n_paragraphs_consumed]
            except IndexError:
                break
            n_paragraphs_consumed += 1

            sty = (p.style.name if p.style else 'Normal')
            txt = p.text.strip()

            # Detectar imágenes en el párrafo (inline o anchored)
            drawing_elems = p._p.findall('.//' + qn('w:drawing'))
            pict_elems = p._p.findall('.//' + qn('w:pict'))

            if drawing_elems or pict_elems:
                # Buscar el rId de la imagen
                for blip in p._p.iter():
                    if blip.tag.endswith('}blip'):
                        rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if rid and rid in rels_map:
                            img_file = rels_map[rid]
                            if img_file in image_files:
                                img_path = image_files[img_file]
                                add_image(dst, img_path, width_in=6.5)
                                n_images_rendered += 1
                                break
                continue  # No procesar más este párrafo

            if not txt and sty in ('Normal', 'CV_Footnote'):
                continue

            if sty == 'Heading 1':
                add_h1(dst, txt)
            elif sty == 'Heading 2':
                add_h2(dst, txt)
            elif sty == 'Heading 3':
                add_h3(dst, txt)
            elif sty == 'Caption':
                add_caption(dst, txt)
            elif sty in ('CV_BulletJustify', 'List Bullet'):
                add_bullet(dst, txt)
            elif sty == 'CV_Footnote':
                add_footnote(dst, txt)
            elif sty == 'Title':
                add_h1(dst, txt)
            else:
                add_paragraph(dst, txt)

        elif tag == 'tbl':
            try:
                tbl = next(table_iter)
            except StopIteration:
                continue

            if is_empty_table(tbl):
                n_tables_skipped += 1
                print(f"  [skip] tabla vacía/degenerada saltada")
                continue

            if is_formula_block_table(tbl):
                # Renderizar como bloque de fórmula (párrafos, sin caja)
                text = tbl.rows[0].cells[0].text
                title, body = formula_title_from_text(text)
                add_formula_block(dst, body, title=title)
                n_callouts += 1
                print(f"  [formula] '{(title or text[:50])}'")
                continue

            # Tabla regular
            headers, rows = extract_table(tbl)
            # Validar headers no vacíos
            if not any(h for h in headers):
                # Si headers están vacíos pero hay rows, usar primera row como header
                if rows:
                    headers = rows[0]
                    rows = rows[1:]

            if not headers or not rows:
                n_tables_skipped += 1
                print(f"  [skip] tabla sin headers o sin filas")
                continue

            widths = col_widths_for(headers, len(rows))
            highlight = detect_highlight(headers, rows)
            add_table_vf3(dst, headers=headers, rows=rows,
                          highlight_row=highlight, col_widths_in=widths)
            n_tables_rendered += 1
            print(f"  [tabla {n_tables_rendered}] {len(headers)}c x {len(rows)}f"
                  + (f" highlight={highlight}" if highlight is not None else ""))

    dst.save(output_path)

    # Limpiar tempdir de imágenes
    try:
        shutil.rmtree(tmpdir)
    except Exception:
        pass

    print(f"\n[OK] Guardado: {output_path}")
    print(f"     Tamaño: {os.path.getsize(output_path)/1024:.1f} KB")
    print(f"     Tablas válidas: {n_tables_rendered}")
    print(f"     Formula blocks: {n_callouts}")
    print(f"     Imágenes:       {n_images_rendered}")
    print(f"     Tablas omitidas (vacías): {n_tables_skipped}")
    return {
        'output_path': output_path,
        'tables': n_tables_rendered,
        'callouts': n_callouts,
        'images': n_images_rendered,
        'skipped': n_tables_skipped,
    }


def derive_output_path(input_path):
    folder = os.path.dirname(input_path)
    base = os.path.splitext(os.path.basename(input_path))[0]
    out = os.path.join(folder, f"{base}_CV.docx")
    if os.path.exists(out):
        i = 2
        while os.path.exists(os.path.join(folder, f"{base}_CV_v{i}.docx")):
            i += 1
        out = os.path.join(folder, f"{base}_CV_v{i}.docx")
    return out


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python reformat_template.py <input.docx> [output.docx]")
        sys.exit(1)
    input_path = sys.argv[1]
    if not os.path.exists(input_path):
        print(f"ERROR: no existe {input_path}")
        sys.exit(1)
    output_path = sys.argv[2] if len(sys.argv) > 2 else derive_output_path(input_path)
    reformat(input_path, output_path)
